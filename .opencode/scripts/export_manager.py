#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
正文导出管理器

支持将章节正文导出为多种格式：
- TXT: 纯文本
- Markdown: Markdown 格式
- EPUB: 电子书
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path
from typing import List, Optional, Tuple


class ExportManager:
    """正文导出管理器"""

    def __init__(self, project_root: str):
        self.project_root = Path(project_root)
        self.novel_dir = self.project_root / "正文"
        self.output_dir = self.project_root / "导出"

    def get_chapter_list(self) -> List[int]:
        """获取章节列表（递归查找子目录）"""
        chapters = []
        if not self.novel_dir.exists():
            return chapters

        # 递归查找所有 .md 文件
        for f in self.novel_dir.rglob("*.md"):
            if f.is_file():
                match = re.search(r"第(\d+)章", f.stem)
                if match:
                    chapters.append(int(match.group(1)))

        return sorted(set(chapters))

    def parse_chapter_range(self, range_str: str) -> List[int]:
        """解析章节范围

        支持格式：
        - "1-10" -> [1, 2, ..., 10]
        - "1,3,5" -> [1, 3, 5]
        - "1,3-5,10" -> [1, 3, 4, 5, 10]
        - "all" -> 全部章节
        """
        if range_str.lower() == "all":
            return self.get_chapter_list()

        chapters = []
        parts = range_str.split(",")

        for part in parts:
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                chapters.extend(range(int(start), int(end) + 1))
            else:
                chapters.append(int(part))

        return sorted(set(chapters))

    def get_chapter_content(self, chapter: int) -> Tuple[str, str]:
        """获取章节内容（递归查找子目录）

        Returns:
            (title, content): 章节标题和正文内容
        """
        padded = f"{chapter:04d}"

        # 递归查找带标题的文件名
        for pattern in [f"**/第{padded}章-*.md", f"**/第{padded}章.md"]:
            matches = list(self.novel_dir.glob(pattern))
            if matches:
                file_path = matches[0]
                content = file_path.read_text(encoding="utf-8")
                # 从文件名提取标题
                title = file_path.stem.replace(f"第{padded}章-", "").replace(f"第{padded}章", "")
                return title or f"第{chapter}章", content

        return f"第{chapter}章", ""

    def export_to_txt(
        self,
        chapters: List[int],
        output_path: str,
        include_title: bool = True,
        add_separator: bool = True,
    ) -> int:
        """导出为 TXT 格式

        Args:
            chapters: 章节列表
            output_path: 输出文件路径
            include_title: 是否包含章节标题
            add_separator: 章节之间是否添加分隔符

        Returns:
            导出的章节数量
        """
        self.output_dir.mkdir(parents=True, exist_ok=True)

        output_file = Path(output_path)
        if output_file.is_dir():
            output_file = self.output_dir / f"{self.project_root.name}.txt"

        with open(output_file, "w", encoding="utf-8") as f:
            for i, chapter in enumerate(chapters):
                title, content = self.get_chapter_content(chapter)

                if include_title:
                    f.write(f"\n{title}\n")
                    f.write("=" * len(title) + "\n\n")

                # 去除 frontmatter（--- 之间的内容）
                lines = content.split("\n")
                in_frontmatter = False
                frontmatter_count = 0

                for line in lines:
                    if line.strip() == "---":
                        frontmatter_count += 1
                        if frontmatter_count == 1:
                            in_frontmatter = True
                            continue
                        elif frontmatter_count == 2:
                            in_frontmatter = False
                            continue
                    if not in_frontmatter and line.strip():
                        f.write(line + "\n")

                if add_separator and i < len(chapters) - 1:
                    f.write("\n" + "=" * 40 + "\n\n")

        print(f"[OK] 已导出 TXT: {output_file} ({len(chapters)} 章)")
        return len(chapters)

    def export_to_markdown(
        self,
        chapters: List[int],
        output_path: str,
    ) -> int:
        """导出为 Markdown 格式

        Args:
            chapters: 章节列表
            output_path: 输出文件路径

        Returns:
            导出的章节数量
        """
        self.output_dir.mkdir(parents=True, exist_ok=True)

        output_file = Path(output_path)
        if output_file.is_dir():
            output_file = self.output_dir / f"{self.project_root.name}.md"

        with open(output_file, "w", encoding="utf-8") as f:
            for i, chapter in enumerate(chapters):
                title, content = self.get_chapter_content(chapter)

                # 章节标题（Markdown 格式）
                f.write(f"\n## {title}\n\n")

                # 去除 frontmatter（--- 之间的内容）
                lines = content.split("\n")
                in_frontmatter = False
                frontmatter_count = 0

                for line in lines:
                    if line.strip() == "---":
                        frontmatter_count += 1
                        if frontmatter_count == 1:
                            in_frontmatter = True
                            continue
                        elif frontmatter_count == 2:
                            in_frontmatter = False
                            continue
                    if not in_frontmatter and line.strip():
                        f.write(line + "\n")

                if i < len(chapters) - 1:
                    f.write("\n---\n")

        print(f"[OK] 已导出 Markdown: {output_file} ({len(chapters)} 章)")
        return len(chapters)

    def export_to_epub(
        self,
        chapters: List[int],
        output_path: str,
    ) -> int:
        """导出为 DOCX 格式"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("错误: 需要安装 python-docx 库")
            print("运行: pip install python-docx")
            return 0

        self.output_dir.mkdir(parents=True, exist_ok=True)

        output_file = Path(output_path)
        if output_file.is_dir():
            output_file = self.output_dir / f"{self.project_root.name}.docx"

        doc = Document()

        # 添加标题
        title = doc.add_heading(self.project_root.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for chapter in chapters:
            chapter_title, content = self.get_chapter_content(chapter)

            # 章节标题
            h = doc.add_heading(chapter_title, level=1)

            # 处理内容
            lines = content.split("\n")
            in_frontmatter = False
            frontmatter_count = 0

            for line in lines:
                if line.strip() == "---":
                    frontmatter_count += 1
                    if frontmatter_count == 1:
                        in_frontmatter = True
                        continue
                    elif frontmatter_count == 2:
                        in_frontmatter = False
                        continue

                if not in_frontmatter and line.strip():
                    p = doc.add_paragraph(line)

        doc.save(str(output_file))
        print(f"[OK] 已导出 DOCX: {output_file} ({len(chapters)} 章)")
        return len(chapters)

    def export_to_epub(
        self,
        chapters: List[int],
        output_path: str,
        author: str = "未知作者",
        language: str = "zh-CN",
    ) -> int:
        """导出为 EPUB 格式"""
        try:
            import ebooklib
            from ebooklib import epub
        except ImportError:
            print("错误: 需要安装 ebooklib 库")
            print("运行: pip install ebooklib")
            return 0

        self.output_dir.mkdir(parents=True, exist_ok=True)

        output_file = Path(output_path)
        if output_file.is_dir():
            output_file = self.output_dir / f"{self.project_root.name}.epub"

        book = epub.EpubBook()

        # 设置元数据
        book.set_identifier(f"novel-{self.project_root.name}")
        book.set_title(self.project_root.name)
        book.set_language(language)
        book.add_author(author)

        spine = ["nav"]
        toc = []

        for chapter in chapters:
            chapter_title, content = self.get_chapter_content(chapter)

            # 创建章节
            c = epub.EpubHtml(
                title=chapter_title,
                file_name=f"chapter_{chapter}.xhtml",
                lang=language,
            )

            # 处理内容为 HTML
            html_content = f"<h1>{chapter_title}</h1>\n"

            lines = content.split("\n")
            in_frontmatter = False
            frontmatter_count = 0

            for line in lines:
                if line.strip() == "---":
                    frontmatter_count += 1
                    if frontmatter_count == 1:
                        in_frontmatter = True
                        continue
                    elif frontmatter_count == 2:
                        in_frontmatter = False
                        continue

                if not in_frontmatter and line.strip():
                    html_content += f"<p>{line}</p>\n"

            c.content = html_content
            book.add_item(c)

            spine.append(c)
            toc.append(c)

        # 设置目录和脊
        book.toc = tuple(toc)
        book.spine = spine

        # 添加导航文件
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # 保存文件
        epub.write_epub(str(output_file), book, {})
        print(f"[OK] 已导出 EPUB: {output_file} ({len(chapters)} 章)")
        return len(chapters)


def main():
    # 启用 Windows UTF-8 支持
    try:
        from runtime_compat import enable_windows_utf8_stdio
        enable_windows_utf8_stdio(skip_in_pytest=True)
    except ImportError:
        pass

    parser = argparse.ArgumentParser(description="正文导出工具")
    parser.add_argument("--project-root", required=True, help="项目根目录")

    sub = parser.add_subparsers(dest="command", required=True)

    # list 命令
    p_list = sub.add_parser("list", help="列出可导出章节")
    p_list.set_defaults(func=cmd_list)

    # export 命令
    p_export = sub.add_parser("export", help="导出正文")
    p_export.add_argument(
        "--range",
        default="all",
        help="章节范围，如 1-10,15,20-30 或 all",
    )
    p_export.add_argument(
        "--format",
        choices=["txt", "markdown", "epub"],
        default="txt",
        help="导出格式",
    )
    p_export.add_argument(
        "--output",
        help="输出文件路径（默认自动生成）",
    )
    p_export.add_argument("--author", default="未知作者", help="作者名（仅 EPUB 需要）")
    p_export.set_defaults(func=cmd_export)

    args = parser.parse_args()

    if args.command == "list":
        manager = ExportManager(args.project_root)
        chapters = manager.get_chapter_list()
        if chapters:
            print(f"可导出章节: {chapters}")
            print(f"共 {len(chapters)} 章")
        else:
            print("未找到章节文件")
        return

    if args.command == "export":
        manager = ExportManager(args.project_root)
        chapters = manager.parse_chapter_range(args.range)

        if not chapters:
            print("没有可导出的章节")
            return

        output_path = args.output or f"{Path(args.project_root).name}.{args.format}"

        if args.format == "txt":
            manager.export_to_txt(chapters, output_path)
        elif args.format == "markdown":
            manager.export_to_markdown(chapters, output_path)
        elif args.format == "epub":
            manager.export_to_epub(chapters, output_path, author=args.author)


def cmd_list(args):
    pass


def cmd_export(args):
    pass


if __name__ == "__main__":
    main()
